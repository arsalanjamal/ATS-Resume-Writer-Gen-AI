import streamlit as st
from fpdf import FPDF
from docx import Document
import googletrans
from googletrans import Translator
import os
import re
import nltk
from nltk.corpus import stopwords

# Ensure NLTK stopwords are downloaded
nltk.download('stopwords')

# Function to create PDF resume
def generate_pdf(data, file_name):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("Arial", "B", 16)
    pdf.cell(200, 10, txt="Resume", ln=True, align="C")

    # Name and contact info
    pdf.set_font("Arial", "B", 12)
    pdf.cell(200, 10, txt=f"Name: {data['name']}", ln=True)
    pdf.cell(200, 10, txt=f"Email: {data['email']}", ln=True)
    pdf.cell(200, 10, txt=f"Phone: {data['phone']}", ln=True)
    pdf.cell(200, 10, txt=f"LinkedIn: {data['linkedin']}", ln=True)

    # Education
    pdf.set_font("Arial", "B", 12)
    pdf.cell(200, 10, txt="Education", ln=True)
    pdf.set_font("Arial", "", 12)
    for education in data['education']:
        pdf.multi_cell(200, 10, txt=f"{education['degree']} - {education['school']} ({education['year']})")

    # Work Experience
    pdf.set_font("Arial", "B", 12)
    pdf.cell(200, 10, txt="Work Experience", ln=True)
    pdf.set_font("Arial", "", 12)
    for experience in data['experience']:
        pdf.multi_cell(200, 10, txt=f"{experience['job_title']} at {experience['company']} ({experience['years']})")

    # Save PDF to file
    pdf.output(file_name)

# Function to create Word resume
def generate_word(data, file_name):
    doc = Document()
    doc.add_heading('Resume', 0)

    doc.add_heading('Name:', level=1)
    doc.add_paragraph(data['name'])

    doc.add_heading('Email:', level=1)
    doc.add_paragraph(data['email'])

    doc.add_heading('Phone:', level=1)
    doc.add_paragraph(data['phone'])

    doc.add_heading('LinkedIn:', level=1)
    doc.add_paragraph(data['linkedin'])

    # Education
    doc.add_heading('Education', level=1)
    for education in data['education']:
        doc.add_paragraph(f"{education['degree']} - {education['school']} ({education['year']})")

    # Work Experience
    doc.add_heading('Work Experience', level=1)
    for experience in data['experience']:
        doc.add_paragraph(f"{experience['job_title']} at {experience['company']} ({experience['years']})")

    doc.save(file_name)

# Function to extract text from uploaded resume
def extract_text_from_file(uploaded_file):
    file_type = uploaded_file.name.split('.')[-1]
    if file_type == 'docx':
        doc = Document(uploaded_file)
        text = '\n'.join([para.text for para in doc.paragraphs])
    else:
        text = uploaded_file.getvalue().decode('utf-8')
    return text

# Function to translate text to selected language
def translate_text(text, target_language):
    translator = Translator()
    translated = translator.translate(text, dest=target_language)
    return translated.text

# Function to calculate ATS Score (Basic Keyword Matching)
def ats_score(resume_text, job_desc_text):
    stop_words = set(stopwords.words("english"))
    resume_words = [word.lower() for word in re.findall(r'\w+', resume_text) if word.lower() not in stop_words]
    job_desc_words = [word.lower() for word in re.findall(r'\w+', job_desc_text) if word.lower() not in stop_words]

    matching_words = set(resume_words) & set(job_desc_words)
    score = len(matching_words) / len(set(job_desc_words)) * 100

    return score

# Language options
LANGUAGES = googletrans.LANGUAGES

# Streamlit App
def main():
    st.title("ATS Resume Maker")

    # Language selection
    st.sidebar.header("Language Selection")
    language = st.sidebar.selectbox("Select your language for the resume:", list(LANGUAGES.values()))

    # Step 1: Manually Enter Resume Data
    st.header("Enter Your Resume Details")

    name = st.text_input("Full Name")
    email = st.text_input("Email")
    phone = st.text_input("Phone")
    linkedin = st.text_input("LinkedIn Profile")

    # Education Section
    st.subheader("Education")
    education_degree = st.text_input("Degree")
    education_school = st.text_input("School/University")
    education_year = st.text_input("Year of Graduation")
    education_data = []
    if st.button("Add Education"):
        education_data.append({
            "degree": education_degree,
            "school": education_school,
            "year": education_year
        })

    # Work Experience Section
    st.subheader("Work Experience")
    work_title = st.text_input("Job Title")
    work_company = st.text_input("Company")
    work_years = st.text_input("Years Worked")
    experience_data = []
    if st.button("Add Work Experience"):
        experience_data.append({
            "job_title": work_title,
            "company": work_company,
            "years": work_years
        })

    # Job Description (For Optimization)
    job_desc = st.text_area("Paste Job Description here (Optional for optimization)")

    # Resume generation button
    if st.button("Generate Resume"):
        if name and email and phone:
            data = {
                'name': name,
                'email': email,
                'phone': phone,
                'linkedin': linkedin,
                'education': education_data,
                'experience': experience_data
            }

            # Generate Resume in PDF and Word
            pdf_file = f"{name}_resume.pdf"
            word_file = f"{name}_resume.docx"
            generate_pdf(data, pdf_file)
            generate_word(data, word_file)

            # Download buttons
            st.download_button(label="Download PDF Resume", data=open(pdf_file, "rb").read(), file_name=pdf_file, mime="application/pdf")
            st.download_button(label="Download Word Resume", data=open(word_file, "rb").read(), file_name=word_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            # Translate resume (if requested)
            if language:
                translated_resume = translate_text("Resume text here...", language)  # Replace with actual resume text
                translated_file = f"{name}_translated_resume_{language}.pdf"
                generate_pdf({'name': name, 'email': email, 'phone': phone, 'linkedin': linkedin}, translated_file)
                st.download_button(label="Download Translated Resume", data=open(translated_file, "rb").read(), file_name=translated_file, mime="application/pdf")

    # Step 2: Upload Existing Resume
    st.header("Or Upload Existing Resume")
    uploaded_file = st.file_uploader("Upload a resume in DOCX format", type=["docx"])
    if uploaded_file:
        resume_text = extract_text_from_file(uploaded_file)
        st.write("Uploaded Resume Content:")
        st.text_area("Resume Text", resume_text, height=200)

    # ATS Score Checker
    st.header("ATS Score Checker")
    if uploaded_file and job_desc:
        score = ats_score(resume_text, job_desc)
        st.write(f"ATS Score: {score}%")

if __name__ == "__main__":
    main()
